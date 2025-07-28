from docx import Document
import subprocess
import os

class VideoProcessor:
    def __init__(self):
        self.temp_dir = "data/temp"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def convert_to_doc(self, video_path):
        """Convert video to DOC with thumbnail"""
        try:
            # Generate thumbnail
            thumb_path = os.path.join(self.temp_dir, f"thumb_{os.path.basename(video_path)}.jpg")
            subprocess.run([
                'ffmpeg', '-i', video_path,
                '-ss', '00:00:05', '-vframes', '1', thumb_path
            ], check=True)
            
            # Create DOC
            doc_path = os.path.splitext(video_path)[0] + '.docx'
            doc = Document()
            doc.add_heading('Video Content', level=1)
            doc.add_picture(thumb_path, width=docx.shared.Inches(4.5))
            doc.add_paragraph("Video will be sent separately in Telegram")
            doc.save(doc_path)
            
            # Cleanup
            os.remove(thumb_path)
            return doc_path
            
        except Exception as e:
            print(f"DOC conversion failed: {str(e)}")
            return None
